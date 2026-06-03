"""
Shared utility functions for AI Humanizer routes.

Includes: DB connection, text extraction, file generation,
analysis helpers, auth decorator, and background task functions.
"""

import io
import os
import logging
from functools import wraps
from flask import g, session, jsonify, send_file


# ========== Database Connection ==========

def get_db():
    """Get a database connection scoped to the current request context."""
    if 'db_conn' not in g:
        from app.models import get_connection
        g.db_conn = get_connection()
    return g.db_conn


def close_db(exception=None):
    """Close the database connection at the end of each request."""
    conn = g.pop('db_conn', None)
    if conn is not None:
        conn.close()


# ========== Auth Decorator ==========

def login_required(f):
    """Require user to be logged in. Returns 401 with login_required flag."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('user_id'):
            return jsonify({"error": "请先登录", "login_required": True}), 401
        return f(*args, **kwargs)
    return decorated_function


# ========== Analysis Helpers ==========

def derive_risk_level(ai_score):
    """Derive risk level from AI score without running full analysis."""
    if not ai_score:
        return "Unknown"
    if ai_score < 20:
        return "Safe"
    elif ai_score < 40:
        return "Warning"
    elif ai_score < 60:
        return "Moderate Risk"
    else:
        return "High Risk"


def generate_modification_suggestions(analysis_result, text):
    """Generate suggestions based on AI analysis results."""
    suggestions = []
    sub_scores = analysis_result.get("sub_scores", {})
    sub_details = analysis_result.get("sub_score_details", {})

    # 1. Perplexity suggestion
    if sub_scores.get("perplexity_score", 0) > 50:
        suggestions.append({
            "target": "perplexity",
            "icon": "📊",
            "title": "词汇多样性不足",
            "detail": "你的文本词汇模式过于可预测，AI检测模型容易识别。建议增加同义词替换和句式变化。",
            "severity": "high" if sub_scores["perplexity_score"] > 70 else "medium"
        })

    # 2. Pattern suggestion
    pattern_data = sub_details.get("pattern", {})
    if pattern_data.get("ai_phrase_count", 0) > 3:
        top_phrases = pattern_data.get("top_phrases", [])
        suggestions.append({
            "target": "pattern",
            "icon": "🔍",
            "title": f"检测到 {pattern_data['ai_phrase_count']} 个AI常用短语",
            "detail": f"常见AI短语如「{'」、「'.join(top_phrases[:3])}」在AI生成文本中频繁出现，替换为更自然的表达可降低AI率。",
            "severity": "high"
        })

    # 3. Readability suggestion
    readability = sub_details.get("readability", {})
    fk_grade = readability.get("flesch_kincaid", 10)
    avg_sent = readability.get("avg_sentence_length", 20)
    if fk_grade > 14 or avg_sent > 25:
        suggestions.append({
            "target": "readability",
            "icon": "✂️",
            "title": f"句长过于均匀（平均 {avg_sent:.0f} 词/句）",
            "detail": "AI生成的文本句子长度变化较小，缺乏人类写作的自然节奏感。建议混合长短句，增加句长变化。",
            "severity": "high" if avg_sent > 30 else "medium"
        })

    # 4. Burstiness suggestion
    if sub_scores.get("burstiness_score", 50) < 30:
        suggestions.append({
            "target": "burstiness",
            "icon": "📏",
            "title": "句式变化不足",
            "detail": "句子长度和结构变化不够丰富。建议混入短句（<10词）和长句（>30词），打破AI写作的规律性。",
            "severity": "medium"
        })

    # 5. Structure suggestion
    structure = sub_details.get("structure", {})
    if structure.get("formulaic_ratio", 0) > 0.2:
        suggestions.append({
            "target": "structure",
            "icon": "🏗️",
            "title": "句式开头较为刻板",
            "detail": "过多句子以「It is」「This is」「There is」等固定模式开头，建议变化句子起始方式。",
            "severity": "medium"
        })

    # Default suggestion if nothing specific
    if not suggestions:
        suggestions.append({
            "target": "general",
            "icon": "✅",
            "title": "文本质量良好",
            "detail": "AI检测指标正常，当前文本不太可能被标记为AI生成。",
            "severity": "low"
        })

    return suggestions


# ========== Text Extraction ==========

def extract_text_from_docx(filepath):
    """Extract text from .docx file."""
    from docx import Document
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return '\n\n'.join(paragraphs)


def extract_text_from_pdf(filepath):
    """Extract text from .pdf file.

    Detects Turnitin reports and skips the first 2 pages automatically.
    """
    import fitz
    doc = fitz.open(filepath)

    # Read first 2 pages to check for Turnitin
    first_two_pages_text = ""
    for i in range(min(2, len(doc))):
        first_two_pages_text += doc[i].get_text()

    is_turnitin = "turnitin" in first_two_pages_text.lower()

    page_count = len(doc)
    text_parts = []
    start_page = 2 if is_turnitin else 0
    for i in range(start_page, page_count):
        text_parts.append(doc[i].get_text())
    doc.close()

    result = '\n\n'.join(text_parts).strip()
    if is_turnitin:
        logging.info(f"Turnitin report detected, skipped first 2 pages ({page_count} pages total)")

    return result


def extract_text(filepath):
    """Extract text from uploaded file based on extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return extract_text_from_docx(filepath)
    elif ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext in ('.txt', '.md'):
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ========== Format Output Helpers ==========

def generate_docx(text):
    """Generate a .docx file in-memory from text content."""
    from docx import Document
    doc = Document()
    for paragraph in text.split('\n\n'):
        p = doc.add_paragraph(paragraph.strip())
        if not paragraph.strip():
            p.add_run('\u200b')  # zero-width space keeps empty paragraph without visible whitespace
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_file_response(text, original_format, filename):
    """Generate a file response for download based on format."""
    base_name = os.path.splitext(filename)[0] if filename else 'humanized'

    # PDF originals have no layout/fonts preserved, so default to docx output
    if original_format == 'pdf':
        original_format = 'docx'

    if original_format == 'docx':
        buf = generate_docx(text)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{base_name}_humanized.docx'
        )
    elif original_format == 'md':
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=f'{base_name}_humanized.md'
        )
    else:  # txt (default)
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f'{base_name}_humanized.txt'
        )


# ========== Background Task Functions ==========

def do_background_rewrite(order_id, text, mode):
    """
    Execute the actual humanization rewrite for a paid order.
    Runs in a background thread. Creates its own DB connection.
    """
    try:
        from app.extensions import humanizer_adapter
        from app.ai_checker import analyze_text
        from app.models import get_connection, Order

        humanized = humanizer_adapter.humanize(text, mode=mode)
        rewritten_analysis = analyze_text(humanized)
        original_analysis = analyze_text(text)

        conn = get_connection()
        try:
            Order.update_result(
                conn, order_id, humanized,
                rewritten_analysis.get('ai_score', 0),
                original_analysis.get('ai_score', 0)
            )
        finally:
            conn.close()
    except Exception:
        logging.exception(f"Background rewrite failed for {order_id}")
        try:
            from app.models import get_connection, Order
            conn = get_connection()
            try:
                Order.mark_failed(conn, order_id)
            finally:
                conn.close()
        except Exception:
            logging.exception(f"Failed to mark order {order_id} as failed")


def process_payment_success(conn, order_id, trade_no):
    """
    Internal function to handle successful payment.
    Marks order as paid and triggers rewrite via thread pool.
    """
    from app.models import Order
    from app.extensions import rewrite_executor

    order = Order.get_by_order_id(conn, order_id)
    if not order:
        logging.error(f"Order {order_id} not found during payment processing")
        return

    # Mark as paid
    from datetime import datetime, timezone
    Order.mark_paid(conn, order_id, trade_no, datetime.now(timezone.utc).isoformat())

    # Read mode from DB
    mode = order.get('mode', 'academic')
    text = order['original_text']

    # Submit rewrite to thread pool (don't block the webhook response)
    rewrite_executor.submit(do_background_rewrite, order_id, text, mode)


def recover_processing_orders():
    """
    Scan for orders stuck in 'processing' status and re-trigger rewrite.
    This handles the case where the server restarted while a rewrite was running.
    """
    try:
        from app.models import get_connection
        from app.extensions import rewrite_executor

        conn = get_connection()
        try:
            cursor = conn.execute("SELECT * FROM orders WHERE status = 'processing'")
            stuck_orders = [dict(row) for row in cursor.fetchall()]
            for order in stuck_orders:
                order_id = order['order_id']
                mode = order.get('mode', 'academic')
                text = order.get('original_text', '')
                if not text:
                    continue
                logging.warning(f"Recovering stuck processing order: {order_id}")
                rewrite_executor.submit(do_background_rewrite, order_id, text, mode)
        finally:
            conn.close()
    except Exception:
        logging.exception("Failed to recover processing orders on startup")
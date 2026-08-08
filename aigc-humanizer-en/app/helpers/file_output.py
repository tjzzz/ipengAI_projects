"""File output helpers: generate downloadable docx / md / txt responses."""

import io
import os
from flask import send_file


def _heading_style_name(level):
    """把标题级别映射为 Word 内置样式名。level 0 为 Title。"""
    if level is None:
        return None
    if level == 0:
        return 'Title'
    return f'Heading {level}'


def _apply_paragraph_style(paragraph, para):
    """根据段落结构标记给段落应用样式（标题/正文/列表/缩进）。"""
    text = para.get('text', '')
    if not text:
        return

    is_heading = para.get('is_heading', False)
    level = para.get('heading_level')

    # 标题样式（Heading 1~9 / Title）
    style_name = _heading_style_name(level)
    if style_name is not None:
        paragraph.style = style_name
        paragraph.add_run(text)
        return

    paragraph.add_run(text)

    # 列表编号：原始编号已还原为文本，直接作为普通文本（保留缩进）
    indent = para.get('indent')
    if indent is not None:
        try:
            # Word 缩进单位 twips（1/20 pt），转 EMU：1 twip = 635 EMU
            paragraph.paragraph_format.left_indent = int(indent) * 635
        except (ValueError, TypeError):
            paragraph.paragraph_format.left_indent = None


def generate_docx(text, paragraphs=None):
    """Generate a .docx file in-memory from text content.

    paragraphs: 可选，结构化段落列表 [{'text','is_heading','heading_level','style',...}]。
        提供时按标题级别重建格式（Title / Heading 1~9 / 正文），否则全部按普通段落。
    """
    from docx import Document
    doc = Document()

    if paragraphs:
        for para in paragraphs:
            ptext = (para.get('text') or '').strip()
            if not ptext:
                continue
            p = doc.add_paragraph()
            _apply_paragraph_style(p, para)
    else:
        for paragraph in text.split('\n\n'):
            p = doc.add_paragraph(paragraph.strip())
            if not paragraph.strip():
                p.add_run('\u200b')  # zero-width space keeps empty paragraph without visible whitespace

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_file_response(text, original_format, filename, paragraphs=None):
    """Generate a file response for download based on format."""
    base_name = os.path.splitext(filename)[0] if filename else 'humanized'

    # PDF originals have no layout/fonts preserved, so default to docx output
    if original_format == 'pdf':
        original_format = 'docx'

    if original_format == 'docx':
        buf = generate_docx(text, paragraphs=paragraphs)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{base_name}_humanized.docx'
        )
    elif original_format == 'md':
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=f'{base_name}_humanized.md'
        )
    else:  # txt (default)
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f'{base_name}_humanized.txt'
        )
